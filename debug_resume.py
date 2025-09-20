#!/usr/bin/env python3
"""
测试简历解析 - 专门分析戴小菲的简历
"""
import sys
import os
import codecs

# 设置编码
if sys.stdout.encoding != 'utf-8':
    sys.stdout = codecs.getwriter('utf-8')(sys.stdout.buffer, 'ignore')
if sys.stderr.encoding != 'utf-8':
    sys.stderr = codecs.getwriter('utf-8')(sys.stderr.buffer, 'ignore')

sys.path.append(os.path.join(os.path.dirname(__file__), 'backend'))

from docx import Document
import re

def analyze_resume_structure(file_path):
    """详细分析简历结构"""
    print(f"[调试] 分析简历文件: {file_path}")

    if not os.path.exists(file_path):
        print(f"[错误] 文件不存在: {file_path}")
        return

    try:
        # 读取Word文档
        doc = Document(file_path)

        print("\n[文档] 段落内容：")
        print("=" * 50)

        for i, paragraph in enumerate(doc.paragraphs):
            text = paragraph.text.strip()
            if text:
                print(f"段落{i+1:2d}: {text}")

        print("\n" + "=" * 50)

        # 合并所有文本
        full_text = '\n'.join([p.text.strip() for p in doc.paragraphs if p.text.strip()])

        print(f"\n[统计] 总文本长度: {len(full_text)} 字符")
        print(f"\n[内容] 完整文本内容：")
        print("-" * 30)
        print(full_text)
        print("-" * 30)

        # 分析可能的字段
        print("\n[分析] 字段分析：")
        analyze_fields(full_text)

        return full_text

    except Exception as e:
        print(f"[错误] 分析失败: {str(e)}")
        import traceback
        traceback.print_exc()
        return None

def analyze_fields(text):
    """分析文本中的各种字段"""
    lines = text.split('\n')

    print("[分析] 逐行分析：")
    for i, line in enumerate(lines):
        if line.strip():
            print(f"  第{i+1}行: {line}")

            # 检查可能的个人信息
            if any(keyword in line for keyword in ['姓名', '电话', '邮箱', '年龄', '性别']):
                print(f"    -> 可能是个人信息")

            # 检查邮箱
            email_pattern = r'\b[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Z|a-z]{2,}\b'
            if re.search(email_pattern, line):
                print(f"    -> 包含邮箱地址")

            # 检查电话
            phone_pattern = r'\d{3,4}[-\s]?\d{3,4}[-\s]?\d{4}'
            if re.search(phone_pattern, line):
                print(f"    -> 包含电话号码")

            # 检查教育背景
            if any(keyword in line for keyword in ['大学', '学院', '学校', '专业', '学历']):
                print(f"    -> 可能是教育背景")

            # 检查工作经历
            if any(keyword in line for keyword in ['公司', '职位', '工作', '经验', '负责']):
                print(f"    -> 可能是工作经历")

            # 检查技能
            if any(keyword in line for keyword in ['技能', '掌握', '熟悉', '精通']):
                print(f"    -> 可能是技能描述")

def improved_personal_info_extraction(text):
    """改进的个人信息提取"""
    print("\n[测试] 改进的个人信息提取测试：")
    personal_info = {}
    lines = text.split('\n')

    # 更灵活的提取逻辑
    for line in lines:
        line = line.strip()
        if not line:
            continue

        # 提取姓名 - 多种模式
        name_patterns = [
            r'姓\s*名[：:]\s*([^\s，,]+)',
            r'([^\s\d]{2,4})\s*求职意向',
            r'^([^\s\d]{2,4})\s*$',  # 单独一行的姓名
        ]

        for pattern in name_patterns:
            match = re.search(pattern, line)
            if match and '姓名' not in personal_info:
                personal_info['姓名'] = match.group(1)
                print(f"  [成功] 提取姓名: {match.group(1)}")
                break

        # 提取求职意向
        if '求职意向' in line:
            job_match = re.search(r'求职意向[：:]?\s*([^\s]+)', line)
            if job_match:
                personal_info['求职意向'] = job_match.group(1)
                print(f"  [成功] 提取求职意向: {job_match.group(1)}")

        # 提取邮箱
        email_pattern = r'\b([A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Z|a-z]{2,})\b'
        email_match = re.search(email_pattern, line)
        if email_match:
            personal_info['邮箱'] = email_match.group(1)
            print(f"  [成功] 提取邮箱: {email_match.group(1)}")

        # 提取电话
        phone_patterns = [
            r'电话[：:]\s*(\d{3,4}[-\s]?\d{3,4}[-\s]?\d{4})',
            r'手机[：:]\s*(\d{11})',
            r'联系方式[：:]\s*(\d{3,4}[-\s]?\d{3,4}[-\s]?\d{4})',
            r'(\d{3}[-\s]?\d{4}[-\s]?\d{4})',  # 通用电话格式
        ]

        for pattern in phone_patterns:
            phone_match = re.search(pattern, line)
            if phone_match:
                personal_info['电话'] = phone_match.group(1)
                print(f"  [成功] 提取电话: {phone_match.group(1)}")
                break

    print(f"\n[结果] 提取结果: {personal_info}")
    return personal_info

if __name__ == "__main__":
    resume_file = r"C:\Users\lyp\Desktop\戴小菲求职意向.docx"

    # 分析简历结构
    full_text = analyze_resume_structure(resume_file)

    # 如果能读取到文本，测试改进的提取逻辑
    if full_text:
        improved_personal_info_extraction(full_text)