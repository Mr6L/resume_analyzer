#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
创建测试用的Word简历文档
"""

from docx import Document

def create_test_resume():
    """创建测试简历文档"""
    doc = Document()

    # 标题
    title = doc.add_heading('个人简历', 0)
    title.alignment = 1  # 居中

    # 个人信息
    doc.add_heading('个人信息', level=1)
    personal_info = [
        "姓名：张三",
        "性别：男",
        "年龄：28",
        "电话：13812345678",
        "邮箱：zhangsan@example.com",
        "地址：北京市朝阳区"
    ]

    for info in personal_info:
        doc.add_paragraph(info)

    # 教育背景
    doc.add_heading('教育背景', level=1)
    doc.add_paragraph('2016-2020 北京大学 计算机科学与技术 本科')
    doc.add_paragraph('主修课程：数据结构、算法设计、计算机网络、数据库原理')

    # 工作经历
    doc.add_heading('工作经历', level=1)
    doc.add_paragraph('2020.7-2023.3 腾讯科技有限公司 软件工程师')
    doc.add_paragraph('负责微信小程序后端开发，使用Python和Django框架')
    doc.add_paragraph('参与用户系统设计，处理日活跃用户超过100万的系统')
    doc.add_paragraph('')
    doc.add_paragraph('2023.4-至今 字节跳动科技有限公司 高级软件工程师')
    doc.add_paragraph('负责抖音推荐算法优化，使用Python和机器学习技术')
    doc.add_paragraph('负责大数据处理和分析，处理PB级别数据')

    # 技能
    doc.add_heading('技能', level=1)
    skills = [
        "Python编程语言，熟练掌握Django、Flask框架",
        "Java编程语言，了解Spring框架",
        "MySQL数据库设计和优化",
        "机器学习算法，熟悉TensorFlow和PyTorch",
        "大数据处理，熟悉Hadoop和Spark"
    ]

    for skill in skills:
        doc.add_paragraph(skill, style='List Bullet')

    # 项目经历
    doc.add_heading('项目经历', level=1)
    doc.add_paragraph('微信小程序用户管理系统')
    doc.add_paragraph('使用Django开发，支持百万级用户注册登录')
    doc.add_paragraph('实现了用户画像分析和个性化推荐功能')
    doc.add_paragraph('')
    doc.add_paragraph('抖音推荐系统优化项目')
    doc.add_paragraph('基于深度学习的个性化推荐算法')
    doc.add_paragraph('提升用户停留时间15%，点击率提升12%')

    # 保存文档
    doc.save('test_resume.docx')
    print("测试简历文档创建成功: test_resume.docx")

if __name__ == "__main__":
    create_test_resume()